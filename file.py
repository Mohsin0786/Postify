from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()

# Add title with center alignment
title = doc.add_heading('Postify API Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add description
doc.add_paragraph('Postify is a Post and Tag Management System API that provides endpoints to manage posts and tags.')

# Base URL
doc.add_heading('Base URL', level=2)
doc.add_paragraph('http://localhost:8000')

# Health Check Endpoint
doc.add_heading('Health Check', level=2)
doc.add_paragraph('Endpoint: GET /health')
doc.add_paragraph('Description: Check if the server is running')
doc.add_heading('Response:', level=3)
doc.add_paragraph('Server is Healthy!!!!')

# Tags Endpoints
doc.add_heading('Tags', level=2)

# Get Tags
doc.add_heading('Get All Tags', level=3)
doc.add_paragraph('Endpoint: GET /api/tags')
doc.add_paragraph('Description: Get tags with pagination and search functionality')
doc.add_paragraph('Query Parameters:')
params = doc.add_paragraph()
params.add_run('• page: Page number (min: 1)\n')
params.add_run('• limit: Items per page (min: 1, max: 100)\n')
params.add_run('• search: Search tags by name (case-insensitive)')

doc.add_heading('Example Response:', level=4)
doc.add_paragraph('''
{
    "success": true,
    "data": [
        {
            "_id": "677cf00500cb0f2087c155f5",
            "name": "frontend"
        }
    ],
    "pagination": {
        "total": 1,
        "page": 1,
        "limit": 10,
        "pages": 1
    }
}
''')

# Create Tag
doc.add_heading('Create Tag', level=3)
doc.add_paragraph('Endpoint: POST /api/tags')
doc.add_paragraph('Description: Create a new tag')
doc.add_paragraph('Request Body:')
doc.add_paragraph('{\n    "name": "sample-tag"\n}')
doc.add_paragraph('Validation Rules:')
rules = doc.add_paragraph()
rules.add_run('• name: Required, 2-30 characters\n')
rules.add_run('• Must be unique (case-insensitive)')

doc.add_heading('Example Response:', level=4)
doc.add_paragraph('''
{
    "success": true,
    "data": {
        "_id": "677cf7620cbc305e26139ae6",
        "name": "sample-tag"
    }
}
''')

# Posts Endpoints
doc.add_heading('Posts', level=2)

# Get Posts
doc.add_heading('Get All Posts', level=3)
doc.add_paragraph('Endpoint: GET /api/posts')
doc.add_paragraph('Description: Get posts with filtering, sorting and pagination')
doc.add_paragraph('Query Parameters:')
post_params = doc.add_paragraph()
post_params.add_run('• sort: Sort by field (allowed: createdAt, title, desc)\n')
post_params.add_run('• page: Page number (min: 1)\n')
post_params.add_run('• limit: Items per page (min: 1, max: 50)\n')
post_params.add_run('• keyword: Search in title and description\n')
post_params.add_run('• tag: Filter by tag IDs (comma-separated for multiple)')

doc.add_heading('Example Response:', level=4)
doc.add_paragraph('''
{
    "success": true,
    "data": [
        {
            "_id": "677cf2ee00cb0f2087c1561b",
            "title": "Sample Post",
            "desc": "Sample description",
            "image": "https://s3.bucket.url/image.jpg",
            "tags": [
                {
                    "_id": "677cf00500cb0f2087c155f5",
                    "name": "frontend"
                }
            ]
        }
    ],
    "pagination": {
        "total": 1,
        "page": 1,
        "limit": 10,
        "pages": 1
    }
}
''')

# Create Post
doc.add_heading('Create Post', level=3)
doc.add_paragraph('Endpoint: POST /api/posts')
doc.add_paragraph('Description: Create a new post with image upload')
doc.add_paragraph('Request Body (multipart/form-data):')
post_body = doc.add_paragraph()
post_body.add_run('• title: Post title (3-100 characters)\n')
post_body.add_run('• desc: Post description (10-5000 characters)\n')
post_body.add_run('• tags: Array of tag IDs (max 10 tags)\n')
post_body.add_run('• image: Image file (JPG, JPEG, PNG, WEBP only, max 5MB)')

doc.add_heading('Example Response:', level=4)
doc.add_paragraph('''
{
    "success": true,
    "data": {
        "title": "Sample Post",
        "desc": "Sample description",
        "tags": ["677cf00500cb0f2087c155f5"],
        "image": "https://s3.bucket.url/image.jpg"
    },
    "message": "Post created successfully"
}
''')

# Save the document
try:
    doc.save('API_Documentation.docx')
    print("Documentation has been generated successfully!")
except Exception as e:
    print(f"Error saving document: {e}") 